# IMPORTS
# =======

from os                     import listdir, path, makedirs
from shutil                 import copyfile, rmtree
from logging                import getLogger, INFO, DEBUG
from argparse               import ArgumentParser
from docx                   import Document
from docx.enum.style        import WD_STYLE_TYPE
from docx.shared            import Pt
from docx.text.paragraph    import Paragraph
from docx.enum.text         import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table        import WD_TABLE_ALIGNMENT
from docx.oxml.ns           import qn
from docx.oxml              import OxmlElement
from re                     import search, sub, findall, match
from datetime               import datetime

import re

# STATIC VARIABLES
# ================

STATIC_DIR          = path.join(path.dirname(path.abspath(__file__)), 'static')
OUTPUT_DIR          = path.join(path.dirname(path.abspath(__file__)), 'output')
EMPTY               = 'EMPTY_LINE'
NOINDENT            = 'NOINDENT'
TOINN               = 'TOINN'    
NUMBER_PREFIX       = '--- '
DEFAULT_GRADE       = 8
PUNCTUATION         = ['.', '?', ':', ';', '!', '…']
HYPERLINK_SCHEMA    = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink'
FONT_NAME           = 'Times New Roman'
FONT_SIZE           = 12

# FUNCTIONS
# =========

def insert_paragraph_before(paragraph, text=''):
    # Get the parent element (usually the body of the document)
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text

    for run in new_para.runs:
        font = run.font
        font.name = FONT_NAME
        font.size = Pt(FONT_SIZE)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    return new_para

def insert_paragraph_after(paragraph, text=''):
    # Get the parent element (usually the body of the document)
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text

    for run in new_para.runs:
        font = run.font
        font.name = FONT_NAME 
        font.size = Pt(FONT_SIZE)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    return new_para

def get_previous_paragraph(paragraph):
    previous_elem = paragraph._p.getprevious()
    while previous_elem is not None:
        if previous_elem.tag.endswith('}p'):  # Check if it's a paragraph element
            return Paragraph(previous_elem, paragraph._parent)
        previous_elem = previous_elem.getprevious()
    return None

def get_next_paragraph(paragraph):
    next_elem = paragraph._p.getnext()
    while next_elem is not None:
        if next_elem.tag.endswith('}p'):  # Check if it's a paragraph element
            return Paragraph(next_elem, paragraph._parent)
        next_elem = next_elem.getnext()
    return None

def starts_with_numbered_prefix(text):
    return bool(re.match(r'^\d+\. ', text))

def remove_paragraph(paragraph):
    p = paragraph._p
    p.getparent().remove(p)
    p._element = None  # Clean up reference

def clean_document(document, args, logger):
    grade = args.grade          if args.grade else DEFAULT_GRADE
    FONT_NAME = 'Courier New'   if grade <= 7 else 'Verdana'
    FONT_SIZE = 12              if grade <= 7 else 13

    logger.info('Cleaning document...')

    # Update all paragraphs and runs
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph and EMPTY in paragraph.text.strip():
            previous_paragraph = get_previous_paragraph(paragraph)
            if previous_paragraph and previous_paragraph.text.strip() == '':
                remove_paragraph(paragraph)
            else:
                paragraph.text = ' '
        if paragraph and TOINN in paragraph.text.strip():
            paragraph.text = paragraph.text.replace(TOINN, '  ')
        if paragraph and paragraph.text.startswith(NOINDENT):
            paragraph.text = paragraph.text.replace(NOINDENT, '')
        if paragraph and paragraph.text.strip() == '':
            if ((previous_paragraph := get_previous_paragraph(paragraph)) and
                (previous_paragraph.text.strip() == '' or previous_paragraph.text.startswith(NUMBER_PREFIX))):
                remove_paragraph(paragraph)
            paragraph.text = ''
        if paragraph and paragraph.style and ('Heading' in paragraph.style.name and
            (previous_paragraph := get_previous_paragraph(paragraph)) and
            not (previous_paragraph.text.strip() == '' or previous_paragraph.text.startswith(NUMBER_PREFIX))):
            insert_paragraph_before(paragraph, ' ')
        if paragraph and (paragraph.text.startswith(NUMBER_PREFIX) and
            (next_paragraph := get_next_paragraph(paragraph)) and
            next_paragraph and next_paragraph.style and
            not 'Heading' in next_paragraph.style.name and
            (previous_paragraph := get_previous_paragraph(paragraph)) and
            not previous_paragraph.text.strip() == ''):
            remove_paragraph(previous_paragraph)

    # =============
    return document

# MAIN
# ====

def main():
    # Parse command line arguments
    parser = ArgumentParser(description='''
        This script cleans a docx file.
        ''')

    parser.add_argument('input',
                        help = 'The input folder')
    parser.add_argument('-o',
                        '--output',
                        help = 'The output epub folder')
    parser.add_argument('-m',
                        '--mathematics',
                        help = 'The epub is a mathematics book',
                        action = 'store_true')
    parser.add_argument('-s',
                        '--science',
                        help = 'The epub is a sience book',
                        action = 'store_true')
    parser.add_argument('-v',
                       '--verbose',
                       help = 'Increase output verbosity',
                       action = 'store_true')
    parser.add_argument('-t',
                        '--toc-levels',
                        help = 'The number of levels in the table of contents',
                        type = int)
    parser.add_argument('-g',
                        '--grade',
                        help = 'The grade level of the book',
                        type = int)
    parser.add_argument('-e',
                        '--excel_creation',
                        help = 'Create Excel files for numeric-heavy tables',
                        action = 'store_true')
    parser.add_argument('-i',
                        '--index',
                        help = 'Do not remove indexes',
                        action = 'store_true')

    args = parser.parse_args()

    # Set up logger
    logger = getLogger(__name__)

    # Set log level
    if args.verbose:
        logger.setLevel(DEBUG)
    else:
        logger.setLevel(INFO)

    # Check if input folder exists
    if not path.exists(args.input):
        print(f'Input path does not exist: {args.input}')
        logger.error('Input file does not exist')
        return

    # Check if input folder is a file
    if not path.isfile(args.input):
        print(f'Input path is not a file: {args.input}')
        logger.error('Input path is not a file')
        return

    # Check if the file has a .docx extension
    if not args.input.lower().endswith('.docx'):
        print(f'Input file is not a DOCX file: {args.input}')
        logger.error('Input file is not a DOCX file')
        return False

    # Try opening the DOCX file to verify it is valid
    try:
        print(args.input)
        document = Document(args.input)  # Attempt to open the DOCX file
    except Exception as e:
        print(f'Invalid DOCX file: {e}')
        logger.error(f'Invalid DOCX file: {e}')
        return False

    # Clean the document
    new_document = clean_document(document, args, logger)
    production_number = args.input.split('/')[-1].split('.')[0]

    # Create output folder. Remove old output folder if it exists
    if path.exists(path.join(OUTPUT_DIR, production_number)):
        logger.info('Removing old output folder')
        rmtree(path.join(OUTPUT_DIR, production_number))
    makedirs(path.join(OUTPUT_DIR, production_number))

    # 14.3 Filnavn
    # ------------

    # Get the title of the document
    LANGUAGES = ["Bokmål", "Nynorsk", "BokmÃ¥l"] # TODO: deal with non unicode characters
    title = document.core_properties.title
    if not title:
        for paragraph in document.paragraphs:
            if paragraph.style.name.lower() == "title" and paragraph.text.strip():
                title = paragraph.text.strip()
                break

    pattern = r"\s+(" + "|".join(LANGUAGES) + r")$"
    title = re.sub(pattern, "", title, flags=re.IGNORECASE).strip()

    # Save the cleaned document
    new_document.save(path.join(OUTPUT_DIR, production_number,  f'{title} {path.basename(args.input)}'))

if __name__ == '__main__':
    main()
