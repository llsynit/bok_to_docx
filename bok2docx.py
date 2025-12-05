#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
bok_to_docx.py — én-filskonverterer for XHTML/EPUB → DOCX

Integrert pipeline (tidligere tre filer samlet i én):
  1) Forberedelse (tidl. prepareXdocx.py)
  2) Pandoc-konvertering (tidl. pandoc_wrapper.py)
  3) Rydding/etterbehandling i DOCX (tidl. clean.py)

Bruk
----
CLI:
  python bok_to_docx.py convert INPUT [-o out.docx] [--stdout]
      [--reference-docx ref.docx] [--pandoc-arg ...]
      [--grade N] [-m|--mathematics] [-s|--science] [-t|--toc-levels N]
      [--no-excel] [--log-level INFO|DEBUG|...]

Eksempler:
  # Fil → DOCX
  python bok_to_docx.py convert book.xhtml -o book.docx

  # Mappa til en utpakket EPUB → DOCX (finner hoved-.xhtml automatisk)
  python bok_to_docx.py convert unpacked_epub_dir -o out.docx

  # Stdin → stdout
  cat book.xhtml | python bok_to_docx.py convert - --stdout > out.docx

Som bibliotek:
  from bok_to_docx import xhtml_to_docx
  bytes_out = xhtml_to_docx(xhtml_bytes, "out.docx", reference_docx=Path("static/referenceDoc.docx"))

Avhenger av:
  - pandoc (må finnes i PATH, eller sett PANDOC_BIN)
  - Python-pakker: beautifulsoup4, lxml, pandas, python-docx

© 2025
"""
from __future__ import annotations

import argparse
import io
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional, Sequence

# --- 3.partsbibliotek (må være installert) -----------------------------------
from bs4 import BeautifulSoup, NavigableString
from pandas import DataFrame
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Konstanter / miljø -------------------------------------------------------
DEFAULT_LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
PANDOC_BIN = os.getenv("PANDOC_BIN", "pandoc")

THIS_DIR = Path(__file__).resolve().parent
STATIC_DIR = THIS_DIR / "static"
OUTPUT_DIR = THIS_DIR / "output"

# Marker/”tokens” brukt både i prepare og clean:
EMPTY = "blank_line"
NOINDENT = "NOINDENT"
TOINN = "TOINN"
NUMBER_PREFIX = "--- "

DEFAULT_GRADE = 8  # barnetrinn-regel (fontvalg mm i clean)
FONT_NAME_JUNIOR = "Courier New"
FONT_NAME_SENIOR = "Verdana"
FONT_SIZE_JUNIOR = 12
FONT_SIZE_SENIOR = 13

# Tittel-opprydding:
LANGUAGES_TRAILING = ["Bokmål", "Nynorsk", "BokmÃ¥l"]  # OBS: inkluder mulig ikke-UTF8-rest

# Excel-deteksjon
MAX_TABLE_WIDTH_CHARS = 54  # ca. bredde for å vurdere om tabell bør ut
EXCEL_DIRNAME = "Excel"

# ------------------------------------------------------------------------------
def configure_logging(level: str = DEFAULT_LOG_LEVEL) -> None:
    logging.basicConfig(
        level=getattr(logging, level, logging.INFO),
        format="%(asctime)s | %(levelname)-7s | %(name)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
        handlers=[logging.StreamHandler(sys.stdout)],
    )

logger = logging.getLogger("bok_to_docx")


# ==============================================================================
# 1) PREPARE (tidligere prepareXdocx.py — forkortet/robustifisert)
# ==============================================================================

def get_table_width_chars(table) -> int:
    """Estimér ca. tegnbredde for tabellen basert på antall kolonner."""
    max_cols = 0
    for row in table.find_all("tr"):
        cells = row.find_all(["td", "th"])
        max_cols = max(max_cols, len(cells))
    # ca. 10 tegn pr kolonne – justerbar heuristikk
    return max_cols * 10


def is_excel_like_table(table) -> bool:
    """
    En enkel gjenkjenning av “Excel-lignende” tabell:
      1) øverste venstre celle tom
      2) første rad (unntatt col 0) A, B, C …
      3) første kolonne (unntatt rad 0) 1, 2, 3 …
    """
    if not table or table.name != "table":
        return False

    rows = table.find_all("tr")
    if len(rows) < 2:
        return False

    first_row_cells = rows[0].find_all(["td", "th"])
    first_col_cells = [r.find(["td", "th"]) for r in rows[1:]]

    # (1) tom øvre-venstre?
    if first_row_cells and first_row_cells[0].get_text(strip=True):
        return False

    # (2) A, B, C...
    expected_headers = [chr(ord("A") + i) for i in range(len(first_row_cells)-1)]
    actual_headers = [c.get_text(strip=True) for c in first_row_cells[1:]]
    if actual_headers[:len(expected_headers)] != expected_headers[:len(actual_headers)]:
        return False

    # (3) 1, 2, 3...
    for i, cell in enumerate(first_col_cells, start=1):
        if not cell:
            return False
        if cell.get_text(strip=True) != str(i):
            return False

    return True

@dataclass
class PrepareArgs:
    mathematics: bool = False
    science: bool = False
    verbose: bool = False
    toc_levels: Optional[int] = None
    grade: Optional[int] = None
    no_excel: bool = False

def _prepend_blank_line(soup, element):
    """Sett inn en EMPTY_LINE før 'element' når hensiktsmessig."""
    empty_p = soup.new_tag("p")
    empty_p.string = EMPTY
    element.insert_before(empty_p)


def prepare_soup(soup: BeautifulSoup, args: ArgumentParser) -> BeautifulSoup:
    print(args)
    """
    Forbered HTML/XHTML for Pandoc/DOCX:
      - pagebreak → ‘--- N’ avsnitt
      - blank linje (EMPTY_LINE) foran overskrifter (m.fl.)
      - enkel regel for figurer/bilder → 'Bilde: alt-tekst'
      - listejustering / innrykk med TOINN
      - "Excel-lignende" tabeller flyttes ut til CSV (hvis ikke --no-excel)
    """
    grade = args.grade if args.grade is not None else DEFAULT_GRADE

    # Moved to statpub_to_bok
    '''
    # 4.12 – sidetall (epub:type=pagebreak) → ‘--- <id|title>’
    args.logger.info("4.12 Sidetall")
    for pagebreak in soup(attrs={"epub:type": "pagebreak"}):
        insert_after = pagebreak
        p = soup.new_tag("p")
        if "title" in pagebreak.attrs:
            p.string = f"{NUMBER_PREFIX}{pagebreak['title']}"
        elif "id" in pagebreak.attrs:
            p.string = f"{NUMBER_PREFIX}{pagebreak['id'].split('-')[-1]}"
        else:
            p.string = f"{NUMBER_PREFIX}".rstrip()

        # sett etter nærmeste p hvis pagebreak var inni et avsnitt
        for parent in pagebreak.parents:
            if parent.name == "p":
                insert_after = parent
                break

        insert_after.insert_after(p)
        pagebreak.decompose()

        # flytt opp om vi står rett inni section
        if p.parent and p.parent.name == "section":
            nxt = p.find_next_sibling()
            if nxt and nxt.name in ("p", "section", "div"):
                nxt.insert_before(p)
    '''

    # Moved to statpub_to_bok
    '''
    # 4.14 – blank linje over innholdsfortegnelse/frontmatter toc
    args.logger.info("4.14 Blank linje over innholdsfortegnelse")
    toc = soup.find("section", attrs={"epub:type": ["frontmatter toc", "toc", "frontmatter"]})
    if toc:
        if toc.previous_sibling and getattr(toc.previous_sibling, "name", None):
            _prepend_blank_line(soup, toc)

    # Blank linje over overskrifter, unntatt når overskrifter følger rett etter hverandre
    args.logger.info("4.14 Blank linje over overskrifter")
    for h in soup.find_all(re.compile(r"^h[1-6]$")):
        prev = h.find_previous_sibling()
        if prev and getattr(prev, "name", None) and not re.match(r"^h[1-6]$", prev.name or ""):
            _prepend_blank_line(soup, h)
    '''

    # Bilder/figurer → “Bilde: <alt>”
    args.logger.info("8.x Bilder/figurer")
    for fig in soup.find_all("figure"):
        # Finn img/alt
        alt = None
        img = fig.find("img")
        if img and img.has_attr("alt"):
            alt = img["alt"]
        if img:
            img.decompose()
        # Sett inn “Bilde: …” før figuren
        p = soup.new_tag("p")
        p.string = "Bilde:" + (f" {alt}" if alt else "")
        fig.insert_before(p)
    
    for fig in soup.find_all("figure"):
        fig.decompose()

    for img in soup.find_all("img"):
        # Unngå duplikat hvis allerede håndtert via <figure>
        if not img.find_parent("figure"):
            alt = img.get("alt", "")
            p = soup.new_tag("p")
            p.string = "Bilde:" + (f" {alt}" if alt else "")
            img.insert_before(p)
            img.decompose()

    # Lister: legg inn logiske innrykk ved hjelp av TOINN
    args.logger.info("4.x Lister → innrykk med TOINN")
    for li in soup.find_all("li"):
        depth = sum(1 for p in li.parents if getattr(p, "name", None) in ("ul", "ol"))
        if depth > 0:
            li.insert(0, NavigableString(TOINN * (depth - 1)))

    # Moved to statpub_to_bok
    '''
    # Oppgavetegn “>>>” → blank linje foran og etter
    args.logger.info("4.1 Blank linje rundt oppgaver (>>>)")
    for p in soup.find_all("p"):
        text = p.get_text(strip=True)
        if text.startswith(">>>"):
            prev = p.previous_sibling
            if prev and getattr(prev, "name", None) != "p":
                _prepend_blank_line(soup, p)
            # blank etter
            empty_after = soup.new_tag("p")
            empty_after.string = EMPTY
            p.insert_after(empty_after)
    '''

    # Excel-lignende tabeller: trekk dem ut (om ikke slått av)
    if not args.no_excel:
        args.logger.info("Tabeller: Excel-sjekk og mulig utskilling")
        excel_root = OUTPUT_DIR / EXCEL_DIRNAME
        if excel_root.exists():
            shutil.rmtree(excel_root)
        excel_root.mkdir(parents=True, exist_ok=True)

        tables_to_extract = []
        for t in soup.find_all("table"):
            width_chars = get_table_width_chars(t)
            if is_excel_like_table(t) or width_chars > MAX_TABLE_WIDTH_CHARS:
                tables_to_extract.append(t)

        # erstatt tabellene og lagre CSV-er
        for idx, t in enumerate(tables_to_extract, start=1):
            rows = []
            for row in t.find_all("tr"):
                cells = [c.get_text(strip=True) for c in row.find_all(["td", "th"])]
                rows.append(cells)
            df = DataFrame(rows)
            csv_path = excel_root / f"table_{idx}.csv"
            df.to_csv(csv_path, index=False, header=False)

            repl = soup.new_tag("p")
            repl.string = f"[Tabell flyttet til {EXCEL_DIRNAME}/table_{idx}.csv]"
            t.insert_before(repl)
            t.decompose()

    # Quick fixes
    for blank_line in soup(attrs={'class':EMPTY}):
        blank_line.string = EMPTY



    # Issues for grade < 8
    # <ul>

    if args.grade and args.grade < 8:

        # UORDNEDE LISTER → div + p med "-- "
        for ul in reversed(soup('ul')):
            new_div = soup.new_tag('div')

            for li in ul.find_all('li', recursive=False):
                # dybde: tell både ul og ol som liste-nivå
                depth = sum(1 for parent in li.parents if parent.name in ('ul', 'ol')) - 1
                indent = '  ' * max(depth, 0)
                parts = []

                for child in li.contents:
                    if getattr(child, 'name', None) in ('ul', 'ol'):
                        continue
                    if hasattr(child, 'get_text'):
                        txt = child.get_text(" ", strip=True)
                    else:
                        txt = str(child).strip()
                    if txt:
                        parts.append(txt)

                text = " ".join(parts).strip()
                if not text:
                    continue

                p = soup.new_tag('p')
                p.string = f"{indent}-- {text}"
                new_div.append(p)

            ul.insert_before(new_div)
            ul.decompose()

        # ORDNEDE LISTER → div + p med '1.' / 'a.' osv.
        for ol in reversed(soup('ol')):
            new_div = soup.new_tag('div')
            type_attr = (ol.get('type') or '').lower()  # f.eks. "a" eller ""

            for idx, li in enumerate(ol.find_all('li', recursive=False), start=1):
                depth = sum(1 for parent in li.parents if parent.name in ('ul', 'ol')) - 1
                indent = '  ' * max(depth, 0)
                parts = []

                for child in li.contents:
                    if getattr(child, 'name', None) in ('ul', 'ol'):
                        continue
                    if hasattr(child, 'get_text'):
                        txt = child.get_text(' ', strip=True)
                    else:
                        txt = str(child).strip()
                    if txt:
                        parts.append(txt)

                text = " ".join(parts).strip()
                if not text:
                    continue

                # velg prefiks: bokstav. eller tall.
                if type_attr == 'a':
                    # a., b., c. ...
                    marker = f"{chr(ord('a') + idx - 1)}."
                else:
                    # 1., 2., 3. ...
                    marker = f"{idx}."

                p = soup.new_tag('p')
                p.string = f"{indent}{marker} {text}"
                new_div.append(p)

            ol.insert_before(new_div)
            ol.decompose()


        # <em> and <strong> in excercises
        for span in soup.select('.exercisenumber'):
            # Hent all tekst inni elementet
            raw = span.get_text()
            # Fjern understreker og trim whitespace/linjeskift
            cleaned = raw.replace('_', '').strip()

            # Tøm innholdet og sett ren tekst
            span.clear()
            if cleaned:
                span.append(NavigableString(cleaned))

    # Remove line breaks in headings indended for PEF
    for element in soup(attrs={'class':'braille-heading-break'}):
        element.decompose()
        # FIXME: merge text nodes?

    # Returnér et renset BeautifulSoup-objekt (som XML/XHTML)
    return BeautifulSoup(str(soup), "lxml-xml")

# ==============================================================================
# 2) PANDOC (tidligere pandoc_wrapper.py)
# ==============================================================================

@dataclass
class PandocOptions:
    pandoc_bin: str = PANDOC_BIN
    reference_docx: Optional[Path] = None
    extra_flags: Sequence[str] = ()


def run_pandoc(args, input_path: Path, output_path: Path, opts: PandocOptions) -> None:
    """
    Kjører pandoc for å lage DOCX av XHTML/HTML.
    Søker default etter static/referenceDoc.docx hvis ikke eksplisitt oppgitt.
    """
    ref_doc = opts.reference_docx
    if ref_doc is None:
        candidate = STATIC_DIR / "referenceDoc.docx"
        if candidate.exists():
            ref_doc = candidate

    cmd = [
        opts.pandoc_bin,
        str(input_path),
        "-f", "html",
        "-t", "docx",
        "-o", str(output_path),
        "--standalone",
    ]
    if ref_doc and ref_doc.exists():
        cmd.extend(["--reference-doc", str(ref_doc)])

    # bruker-flagg til slutt (f.eks --strip-comments, --toc)
    if opts.extra_flags:
        cmd.extend(list(opts.extra_flags))

    args.logger.info("Kjører pandoc: %s", " ".join(cmd))
    try:
        subprocess.run(cmd, check=True)
    except FileNotFoundError as e:
        raise RuntimeError(f"Finner ikke pandoc-binæren {opts.pandoc_bin!r}. Sett PANDOC_BIN eller legg i PATH.") from e
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Pandoc feilet med exit {e.returncode}") from e


# ==============================================================================
# 3) CLEAN (tidligere clean.py — kondensert, men dekker hovedregler)
# ==============================================================================

def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
    return new_para


def _get_previous_paragraph(paragraph: Paragraph) -> Optional[Paragraph]:
    prev = paragraph._p.getprevious()
    while prev is not None:
        if prev.tag.endswith("}p"):
            return Paragraph(prev, paragraph._parent)
        prev = prev.getprevious()
    return None


def _remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._element = None


@dataclass
class CleanArgs:
    grade: Optional[int] = None


def clean_document(document: Document, args: CleanArgs) -> Document:
    """
    Rydder i DOCX-innholdet:
      - EMPTY_LINE → blank avsnitt (men unngå doble)
      - TOINN → 2 mellomrom
      - NOINDENT fjernes først i linje
      - blank linje foran overskrifter hvis nødvendig
      - fjern tomme avsnitt etter tallinje '--- N'
    + enkel fontpolicy basert på trinn
    """
    grade = args.grade if args.grade is not None else DEFAULT_GRADE
    font_name = FONT_NAME_JUNIOR if grade <= 7 else FONT_NAME_SENIOR
    font_size = FONT_SIZE_JUNIOR if grade <= 7 else FONT_SIZE_SENIOR

    args.logger.info("Rydder DOCX …")

    for i, p in enumerate(list(document.paragraphs)):
        text = p.text or ""

        # Tokens
        if text.strip() == EMPTY:
            prev = _get_previous_paragraph(p)
            if prev and prev.text.strip() == "":
                _remove_paragraph(p)
                continue
            p.text = " "

        if TOINN in text:
            p.text = text.replace(TOINN, "  ")
            text = p.text

        if text.startswith(NOINDENT):
            p.text = text.replace(NOINDENT, "", 1)
            text = p.text

        # Unngå dobbel blank etter ‘--- N’
        if text.strip() == "":
            prev = _get_previous_paragraph(p)
            if prev and (prev.text.strip() == "" or prev.text.strip().startswith(NUMBER_PREFIX)):
                _remove_paragraph(p)
                continue

        # Blank linje før overskrifter (hvis ikke blank/side nr. fra før)
        if p.style and "Heading" in p.style.name:
            prev = _get_previous_paragraph(p)
            if prev and not (prev.text.strip() == "" or prev.text.strip().startswith(NUMBER_PREFIX)):
                _insert_paragraph_after(prev, "")

        # Sett font for alle runs
        for run in p.runs:
            f = run.font
            f.name = font_name
            f.size = Pt(font_size)
            # Øst-Asia fallback (Word)
            if run._element.rPr is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    return document


def _derive_title(document: Document) -> str:
    """Hent tittel fra core properties eller første avsnitt med stil 'Title'."""
    title = document.core_properties.title or ""
    if not title:
        for p in document.paragraphs:
            try:
                if p.style and p.style.name.lower() == "title" and p.text.strip():
                    title = p.text.strip()
                    break
            except Exception:
                pass

    if title:
        # fjern evt. trailing språk-ord
        pattern = r"\s+(" + "|".join(map(re.escape, LANGUAGES_TRAILING)) + r")$"
        title = re.sub(pattern, "", title, flags=re.IGNORECASE).strip()
    return title or "Dokument"


# ==============================================================================
# 4) HØYNIVÅ-PIPELINE
# ==============================================================================

def _read_soup_from_input(input_path: Path) -> tuple[BeautifulSoup, Path]:
    """
    Les XHTML enten fra en fil eller fra en mappe med utpakket EPUB.
    Returnerer (soup, faktisk_xhtml_fil).
    """
    if input_path.is_dir():
        # Finn første .xhtml som ikke er nav.xhtml
        xhtml_files = [p for p in input_path.glob("*.xhtml") if p.name.lower() != "nav.xhtml"]
        if not xhtml_files:
            raise FileNotFoundError(f"Fant ingen .xhtml i {input_path}")
        if len(xhtml_files) > 1:
            args.logger.warning("Fant flere .xhtml-filer, bruker: %s", xhtml_files[0].name)
        target = xhtml_files[0]
    else:
        target = input_path
        if target.suffix.lower() not in {".xhtml", ".html", ".htm"}:
            args.logger.warning("Input ser ikke ut som XHTML/HTML: %s", target.name)

    with target.open("rb") as f:
        data = f.read()
    # bruk XML-parser, men tåler HTML
    soup = BeautifulSoup(data, "lxml-xml")
    return soup, target


def xhtml_to_docx(
    xhtml_bytes: bytes,
    output_filename: str = "output.docx",
    *,
    reference_docx: Optional[Path] = None,
    pandoc_args: Sequence[str] = (),
    grade: Optional[int] = None,
    mathematics: bool = False,
    science: bool = False,
    toc_levels: Optional[int] = None,
    no_excel: bool = False,
) -> bytes:
    """API: bytes → bytes."""
    prep_args = PrepareArgs(
        mathematics=mathematics,
        science=science,
        verbose=False,
        toc_levels=toc_levels,
        grade=grade,
        no_excel=no_excel,
    )
    pandoc_opts = PandocOptions(reference_docx=reference_docx, extra_flags=pandoc_args)
    clean_args = CleanArgs(grade=grade)

    with tempfile.TemporaryDirectory(prefix="bok_to_docx_") as tmp_s:
        tmp = Path(tmp_s)
        in_html = tmp / "in.xhtml"
        in_html.write_bytes(xhtml_bytes)

        # prepare
        soup = BeautifulSoup(xhtml_bytes, "lxml-xml")
        prepared = prepare_soup(soup, args, args.logger)
        prepared_html = tmp / "prepared.xhtml"
        prepared_html.write_text(str(prepared), encoding="utf-8")

        # pandoc
        out_docx = tmp / Path(output_filename).name
        run_pandoc(prepared_html, out_docx, pandoc_opts)
        if not out_docx.exists():
            raise RuntimeError("Pandoc genererte ingen DOCX")

        # clean
        doc = Document(str(out_docx))
        doc = clean_document(doc, clean_args, args.logger)

        # output as bytes
        final_docx = tmp / Path(output_filename).name
        doc.save(str(final_docx))
        return final_docx.read_bytes()


# New convert method
def convert(args):
    pandoc_opts = PandocOptions(reference_docx=args.reference_docx, extra_flags=args.pandoc_args or ())
    clean_args = CleanArgs(grade=args.grade)

    # --- INPUT: fil, mappe, eller stdin
    if args.input == "-":
        data = sys.stdin.buffer.read()
        if not data:
            args.logger.error("Ingen data på stdin.")
            return 2
        out_name = args.output.name if args.output else "output.docx"
        try:
            out_bytes = xhtml_to_docx(
                data, out_name,
                reference_docx=args.reference_docx,
                pandoc_args=args.pandoc_args or (),
                grade=args.grade,
                mathematics=args.mathematics,
                science=args.science,
                toc_levels=args.toc_levels,
                no_excel=args.no_excel,
            )
        except Exception as e:
            args.logger.error("Konvertering feilet: %s", e)
            return 1

        if args.stdout:
            sys.stdout.buffer.write(out_bytes)
        else:
            if not args.output:
                args.logger.error("Må angi -o/--output hvis ikke --stdout.")
                return 2
            args.output.parent.mkdir(parents=True, exist_ok=True)
            args.output.write_bytes(out_bytes)
            args.logger.info("Skrev %s (%.1f KB)", args.output, len(out_bytes)/1024.0)
        return 0

    # Fil eller mappe
    in_path = Path(args.input).expanduser().resolve()
    if not in_path.exists():
        args.logger.error("Finner ikke input: %s", in_path)
        return 2

    try:
        soup, actual_xhtml = _read_soup_from_input(in_path)
    except Exception as e:
        args.logger.error("Klarte ikke lese input: %s", e)
        return 2

    with tempfile.TemporaryDirectory(prefix="bok_to_docx_") as tmp_s:
        tmp = Path(tmp_s)

        # prepare
        prepared = prepare_soup(soup, args)
        prepared_html = tmp / (actual_xhtml.name if actual_xhtml.suffix else "prepared.xhtml")
        prepared_html.write_text(str(prepared), encoding="utf-8")

        # pandoc
        out_docx_tmp = tmp / ((args.output.name if args.output else actual_xhtml.with_suffix(".docx").name))
        try:
            run_pandoc(args, prepared_html, out_docx_tmp, pandoc_opts)
        except Exception as e:
            args.logger.error("Pandoc feilet: %s", e)
            return 1

        if not out_docx_tmp.exists():
            args.logger.error("Ingen DOCX fra pandoc.")
            return 1

        # clean
        try:
            doc = Document(str(out_docx_tmp))
            doc = clean_document(doc, args)
        except Exception as e:
            args.logger.error("DOCX-etterbehandling feilet: %s", e)
            return 1

        # Bestem endelig output
        if args.stdout:
            bio = io.BytesIO()
            doc.save(bio)
            sys.stdout.buffer.write(bio.getvalue())
            return 0

        '''
        if args.output:
            args.output.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(args.output))
            args.logger.info("Skrev %s", args.output)
            return 0
        '''

        # Default: lag mappen output/<production_number>/, og lag filnavn ut fra tittel
        production_number = actual_xhtml.stem
        '''
        out_dir = OUTPUT_DIR / production_number
        if out_dir.exists():
            args.logger.info("Fjerner gammel output-mappe: %s", out_dir)
            shutil.rmtree(out_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        '''

        if args.job_dir.exists():
            args.logger.info("Fjerner gammel output-mappe: %s", args.job_dir)
            shutil.rmtree(args.job_dir)
        args.job_dir.mkdir(parents=True, exist_ok=True)

        title = _derive_title(doc)
        #final_name = f"{title} {actual_xhtml.with_suffix('.docx').name}"
        #final_path = out_dir / final_name
        final_name = f"{production_number}.docx"
        final_path = args.job_dir / final_name
        doc.save(str(final_path))
        args.logger.info("Skrev %s", final_path)
        return 0

    args.logger.error("Ukjent kommando.")
    return 2

# ==============================================================================
# 5) CLI
# ==============================================================================

# -----------------------------------------------------------------------------
# Logger (for CLI)
# -----------------------------------------------------------------------------

def make_logger() -> logging.Logger:
    """Logger som skriver til stdout (synlig i docker logs)."""
    logger = logging.getLogger("bok_to_docx")
    if not logger.handlers:
        logger.setLevel(logging.INFO)
        h = logging.StreamHandler()
        h.setFormatter(logging.Formatter("[%(levelname)s] %(message)s"))
        logger.addHandler(h)
    return logger

def _build_arg_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="bok_to_docx",
        description="Konverter XHTML/utpakket EPUB → DOCX i én fil (prepare + pandoc + clean).",
    )

    # Generelle flagg
    p.add_argument(
        "--log-level",
        default=DEFAULT_LOG_LEVEL,
        choices=["CRITICAL", "ERROR", "WARNING", "INFO", "DEBUG"],
        help="Loggnivå",
    )
    p.add_argument(
        "--reference-docx",
        type=Path,
        default=None,
        help="Referanse-DOCX (overstyrer static/referenceDoc.docx).",
    )
    p.add_argument(
        "--pandoc-args",
        action="append",
        default=[],
        help="Ekstra flagg til pandoc (kan gjentas).",
    )
    p.add_argument(
        "--stdout",
        action="store_true",
        help="Skriv DOCX til stdout (binær).",
    )
    p.add_argument(
        "--no-excel",
        action="store_true",
        help="Skru av Excel-uttrekk av tabeller.",
    )

    # prepare/clean-relaterte flagg
    p.add_argument(
        "-g", "--grade",
        type=int,
        help="Klassetrinn (påvirker fontvalg i DOCX).",
    )
    p.add_argument(
        "-m", "--mathematics",
        action="store_true",
        help="Matematikkbok (hint).",
    )
    p.add_argument(
        "-s", "--science",
        action="store_true",
        help="Naturfag (hint).",
    )
    p.add_argument(
        "-t", "--toc-levels",
        type=int,
        help="Antall TOC-nivåer (hint til prepare).",
    )

    # Hoved-IO-argumenter (uten "convert")
    p.add_argument(
        "input",
        help="XHTML/HTML-fil, mappe (utpakket EPUB), eller '-' for stdin.",
    )
    p.add_argument(
        "-o", "--output",
        type=Path,
        help="Output .docx",
    )

    return p

def main(argv: Optional[Sequence[str]] = None) -> int:
    args = _build_arg_parser().parse_args(argv)
    args.logger = make_logger()
    result = convert(args)

    match result:
        case 0:
            args.logger.info("Ferdig uten feil.")
        case 1:
            args.logger.error("Ferdig med feil under konvertering.")
        case 2:
            args.logger.error("Ferdig med feil i input eller parametre.")
        case _:
            args.logger.error("Ukjent avslutningskode: %s", result)


if __name__ == "__main__":
    raise SystemExit(main())
